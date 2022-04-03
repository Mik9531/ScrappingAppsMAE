from google_play_scraper import app
import pandas as pd
from docx import Document

result = app(
    'com.nianticlabs.pokemongo',
    lang='en', # defaults to 'en'
    country='us' # defaults to 'us'
)

df = pd.DataFrame(result)


print(result)

# document = Document()
#
# document.add_heading(result.get('title'))
# document.add_paragraph(result.get('description'))
#
# document.save('output.docx')
